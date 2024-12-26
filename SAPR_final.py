from tkinter import *
import sys
import os
import shutil
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side, Alignment
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
from datetime import datetime

# Проверяем введённое значение
def validate_data(data):

    try:
        float(data)

    except ValueError:
        print('Неверный тип данных!')
        return False

    except Exception:
        print('Непредвиденная ошибка!')
        return False
    
    else:
        return True

'''
    else:
        if float(data) < 0:
            print('Введено отрицательное число или ноль!')
            return False

        else:
            return True'''


def convert_to_number(number):

    number = number.replace(',', '.')
    e_index = number.find('e')

    if e_index == -1:
        return float(number)

    return (float(number[:e_index]) * (10 ** int(number[e_index + 1:])))


# Масштабирование для нормального отображения
def convert_to_si(convert_direction):

    global L_si 
    global A_si 
    global q_si 
    global Sigma_si
    L_si = 100

    # Параметр, для масштабной отрисовки стержней
    global width_canv
    global length_conv
    length_conv = sum(rods_data[1]) * L_si / width_canv * n_rod / 2 if sum(rods_data[1]) >= width_canv else 1

    global rods_start

    if convert_direction:

        rods_start = rods_data[1].copy()

    A_si = 3
    q_si = 1000
    Sigma_si = 100_000_000

    for i in range(len(rods_data[1])):

        rods_data[1][i] = rods_data[1][i] / length_conv * L_si if convert_direction else rods_start[i]

    for i in range(len(rods_data[2])):

        rods_data[2][i] = rods_data[2][i] * A_si if convert_direction else rods_data[2][i] / A_si

    for i in range(len(rods_forces)):

        rods_forces[i] = rods_forces[i] * q_si if convert_direction else rods_forces[i] / q_si

    for i in range(len(nodes_forces)):

        nodes_forces[i] = nodes_forces[i] * q_si if convert_direction else nodes_forces[i] / q_si



#---------------------------- Пре-процессор ----------------------------

# Отрисовка сетки
def draw_grid(canvas, x0, y0, width1, height1):

    # Вертикальные линии
    for i in range(x0, width1, 100):
        canvas.create_line(i, 0, i, height1, fill='#5f6690', width=1)

    for i in range(x0, 0, -100):
        canvas.create_line(i, 0, i, height1, fill='#5f6690', width=1)
    
    # Горизонтальные линии
    for i in range(y0, height1, 100):
        canvas.create_line(0, i, width1, i, fill='#5f6690', width=1)

    for i in range(y0, 0, -100):
        canvas.create_line(0, i, width1, i, fill='#5f6690', width=1)


# Ввод данных через консоль
def input_data():

    global n_rod
    global hard_join_end
    global hard_join_start
    global rods_data
    global rods_forces
    global nodes_forces

    # Количество стержней
    while not validate_data(n_rod := input('Введите количество стержней: ')):
        pass

    n_rod = int(n_rod)

    # Жёсткая заделка в начале
    while (not (validate_data(hard_join_start := input('Есть ли жёсткая заделка в начале? (1 - Да; 0 - Нет) -> ')))) or (not (hard_join_start in ['0', '1'])):
        pass

    hard_join_start = int(hard_join_start)

    # Жёсткая заделка в конце
    while (not (validate_data(hard_join_end := input('Есть ли жёсткая заделка в конце? (1 - Да; 0 - Нет) -> ')))) or (not (hard_join_end in ['0', '1'])):
        pass

    hard_join_end = int(hard_join_end)

    rods_data = {1 : [], 2 : [], 3 : [], 4 : []}
    data_change = ['Длина ', 'Площадь поперечного сечения ', 'Модуль упругости ', 'Допускаемое напряжение ']

    # Ввод данных конструкции
    for j in range(len(data_change)):
        for i in range(n_rod):

            while not (validate_data(data := input(f'{data_change[j]} {i + 1}-го стержня: '))):
                pass

            rods_data[j + 1].append(float(data))

    rods_forces = []
    nodes_forces = []

    print('Введите нагрузки на стержни (Если нагрузки нет - 0)')
    for i in range(n_rod):

        while not (validate_data(data := input(f'Продольная нагрузка в стержне {i + 1}: '))):
            pass

        rods_forces.append(float(data))


    print('Введите нагрузки на узлы (Если нагрузки нет - 0)')
    for i in range(n_rod + 1):

        while not (validate_data(data := input(f'Сосредоточенная нагрузка в узле {i + 1}: '))):
            pass

        nodes_forces.append(float(data))


# Отрисовка длин и площадей
def draw_parametrs(canvas, rods_data):

    start = mid_height
    end_x = start
    max_height = max(rods_data[2])

    for i in range(n_rod):

        canvas.create_line(end_x, middle, end_x, middle + (max_height * 2.5), width=2)
        canvas.create_line(end_x, middle + (max_height * 2.5), end_x + rods_data[1][i], middle + (max_height * 2.5), arrow="both", width=2)
        canvas.create_text(end_x + (rods_data[1][i] / 2), (middle - 10) + (max_height * 2.5), text=f"{rods_start[i]}L, {rods_data[2][i] / A_si}A", font=("Times New Roman", 12, "bold"))

        canvas.create_line(end_x + (rods_data[1][i] / 8), middle - (rods_data[2][i] / 4), end_x + (rods_data[1][i] / 8) + 20, middle - (max_height * 2), width=2)
        canvas.create_line(end_x + (rods_data[1][i] / 8) + 20, middle - (max_height * 2), end_x + (rods_data[1][i] / 8) + 50, middle - (max_height * 2), width=2)
        canvas.create_text(end_x + (rods_data[1][i] / 8) + 35, (middle - 10) - (max_height * 2), text=f"{rods_data[3][i]}E, {rods_data[4][i]}σ", font=("Times New Roman", 12, "bold"))

        end_x += rods_data[1][i]

    canvas.create_line(end_x, middle, end_x, middle + (max_height * 2.5), width=2)
    end_x = start


# Отрисовка сил в стержнях
def draw_forces_rods(canvas, rods_forces, rods_data):
    
    start = mid_height
    end_x = start
    max_height = max(rods_data[2])

    # Отрисовка нагрузок на стержни
    for i in range(n_rod):

        # Продольные нагрузки в стержнях
        if rods_forces[i] > 0:
            canvas.create_line(end_x, middle, end_x + rods_data[1][i], middle, arrow='last', dash=20, width=5, fill='blue')

        elif rods_forces[i] < 0:
            canvas.create_line(end_x, middle, end_x + rods_data[1][i], middle, arrow='first', dash=20, width=5, fill='blue')

        else:
            pass

        if rods_forces[i] != 0:
            canvas.create_line(end_x + (rods_data[1][i] / 2), middle, end_x + (rods_data[1][i] / 2), middle - (max_height * 1.5), width=1, fill='black')
            canvas.create_text(end_x + (rods_data[1][i] / 2), middle - (max_height * 1.7), text=f'{rods_forces[i] / q_si}q', fill='black', font=("Times New Roman", 12, "bold"))

        end_x += rods_data[1][i]

    end_x = start


# Отрисовка сил в узлах
def draw_forces_nodes(canvas, nodes_forces, rods_data):

    start = mid_height
    end_x = start
    max_height = max(rods_data[2])

    # Сосредоточенные нагрузки в узлах
    for i in range(n_rod + 1):

        if nodes_forces[i] != 0:

            canvas.create_line(end_x, middle, end_x + (rods_data[1][i - 1 if i == n_rod  else i] / 4) if nodes_forces[i] > 0 else end_x - (rods_data[1][i - 1 if i == n_rod  else i] / 4), middle, arrow='last', width=5, fill='red')
            canvas.create_line(end_x + (rods_data[1][i - 1 if i == n_rod  else i] / 8) if nodes_forces[i] > 0 else end_x - (rods_data[1][i - 1 if i == n_rod  else i] / 8), middle, end_x + (rods_data[1][i - 1 if i == n_rod  else i] / 8) if nodes_forces[i] > 0 else end_x - (rods_data[1][i - 1 if i == n_rod  else i] / 8), middle - (max_height * 1.5), width=1, fill='black')
            canvas.create_text(end_x + (rods_data[1][i - 1 if i == n_rod  else i] / 8) if nodes_forces[i] > 0 else end_x - (rods_data[1][i - 1 if i == n_rod  else i] / 8), middle - (max_height * 1.7), text=f'{nodes_forces[i] / q_si}qL', fill='black', font=("Times New Roman", 12, "bold"))

        else:
            pass
        
        if i < n_rod:
            end_x += rods_data[1][i]

    end_x = start


# Отрисовка конструкции
def draw_construction(display):

    global L_si
    global q_si

    width = 1100
    height = 600

    global width_canv
    global height_canv
    width_canv = 1200
    height_canv = 800

    convert_to_si(True)

    screen = Tk()
    screen.title('SAPR')
    screen.geometry(f'{width}x{height}')
    screen.iconbitmap(r'logo.ico')

    global canvas
    canvas = Canvas(screen, bg="#aaaaaa", width=width_canv, height=height_canv, scrollregion=(0, -10000, width_canv, 10000))

    hbar=Scrollbar(screen, orient=HORIZONTAL)
    hbar.pack(side=BOTTOM, fill=X)
    hbar.config(command=canvas.xview)

    canvas.config(xscrollcommand=hbar.set)
    canvas.pack()

    global middle
    middle = 250

    global mid_height
    mid_height = 60

    draw_grid(canvas, mid_height, middle, width_canv, height)

    # Отрисовка заделки в начале
    if hard_join_start == True:

        canvas.create_line(mid_height, middle - 50, mid_height, middle + 50, fill='black', width=3)

        canvas.create_line(mid_height, middle - 50, mid_height - 10, middle - 40, fill='black', width=3)
        canvas.create_line(mid_height, middle - 35, mid_height - 10, middle - 25, fill='black', width=3)
        canvas.create_line(mid_height, middle - 20, mid_height - 10, middle - 10, fill='black', width=3)
        canvas.create_line(mid_height, middle - 5, mid_height - 10, middle + 5, fill='black', width=3)
        canvas.create_line(mid_height, middle + 10, mid_height - 10, middle + 20, fill='black', width=3)
        canvas.create_line(mid_height, middle + 25, mid_height - 10, middle + 35, fill='black', width=3)
        canvas.create_line(mid_height, middle + 40, mid_height - 10, middle + 50, fill='black', width=3)

    start = mid_height
    end_x = start
    # Отрисовка стержней
    for i in range(n_rod):

        canvas.create_rectangle(end_x, middle + (rods_data[2][i] / 2), end_x + (rods_data[1][i]), middle - (rods_data[2][i] / 2), width=3)
        end_x += rods_data[1][i]

    # Отрисовка заделки в конце
    if hard_join_end == True:

        canvas.create_line(end_x, middle - 50, end_x, middle + 50, fill='black', width=3)

        canvas.create_line(end_x, middle - 40, end_x + 10, middle - 50, fill='black', width=3)
        canvas.create_line(end_x, middle - 25, end_x + 10, middle - 35, fill='black', width=3)
        canvas.create_line(end_x, middle - 10, end_x + 10, middle - 20, fill='black', width=3)
        canvas.create_line(end_x, middle + 5, end_x + 10, middle - 5, fill='black', width=3)
        canvas.create_line(end_x, middle + 20, end_x + 10, middle + 10, fill='black', width=3)
        canvas.create_line(end_x, middle + 35, end_x + 10, middle + 25, fill='black', width=3)
        canvas.create_line(end_x, middle + 50, end_x + 10, middle + 40, fill='black', width=3)

    end_x = start

    max_height = max(rods_data[2])

    # Указатели стержней
    for i in range(n_rod):

        canvas.create_oval(end_x + ((rods_data[1][i] / 2) - 10), middle - (max_height * 3.5), end_x + ((rods_data[1][i] / 2) + 10), middle - (max_height * 3.5) - 20, fill='#4682B4', outline='black')
        canvas.create_text(end_x + (rods_data[1][i] / 2), middle - (max_height * 3.5) - 10, text=f"{i + 1}", fill="black", font=("Times New Roman", 12, "bold"))

        end_x += rods_data[1][i]

    end_x = start

    if display == 1:
        draw_parametrs(canvas, rods_data)

    elif display == 2:
        draw_forces_rods(canvas, rods_forces, rods_data)

    elif display == 3:
        draw_forces_nodes(canvas, nodes_forces, rods_data)

    elif display in [4, 5, 6]:
        draw_epure(canvas, display)

    # Указатели узлов
    for i in range(n_rod + 1):

        canvas.create_rectangle(end_x - 10, middle + (max_height * 2), end_x + 10, middle + (max_height * 2) + 20, fill='#4682B4', outline='black')
        canvas.create_text(end_x, middle + (max_height * 2) + 10, text=f"{i + 1}", fill="black", font=("Times New Roman", 12, "bold"))

        if i < n_rod:
            end_x += rods_data[1][i]

    end_x = start

    canvas.create_text(950, 60, text=f'Количество стержней: {n_rod}', fill="black", font=("Times New Roman", 14, "bold"))
    canvas.create_text(950, 80, text=f'q = {q_si}', fill="black", font=("Times New Roman", 14, "bold"))

    convert_to_si(False)

    screen.mainloop()


# Считывание с файла
def read_data(file_name=''):

    global n_rod
    global hard_join_end
    global hard_join_start
    global rods_data
    global rods_forces
    global nodes_forces

    n_rod = 0
    hard_join_start = 0
    hard_join_end = 0
    rods_data = {1 : [], 2 : [], 3 : [], 4 : []}
    rods_forces = []
    nodes_forces = []

    if file_name == '':

        file_name = str(input("Введите имя файла: "))

        # Проверка наличия расширения
        if file_name.find('.txt') == -1:
            file_name += '.txt'

    key_words = ['N', 'Hard_join_start', 'Hard_join_end', 'L', 'A', 'E', 'Sigma', 'rods_forces', 'nodes_forces']

    while True:

        try:
            file = open(file_name, 'r')

        except FileNotFoundError:

            try:
                file = open(f"{file_name[:-4]}-folder/{file_name}", 'r')

            except FileNotFoundError:

                print('Такого файла не существует!\n')
                return False
            
            else:
                print(f"Файл открыт в папке проекта: {file_name[:-4]}-folder")

        else:
            print('Файл успешно открыт!')

        for line in file:

            if line.find(key_words[0]) != -1:
                n_rod = int(line[line.find('=') + 1:])

            elif line.find(key_words[1]) != -1:
                hard_join_start = int(line[line.find('=') + 1:])

            elif line.find(key_words[2]) != -1:
                hard_join_end = int(line[line.find('=') + 1:])

            elif line.find(key_words[3]) != -1:
                rods_data[1] = [float(L) for L in line[2:].replace('\n', ' ').replace('\t', ' ').split(' ') if L != '']

            elif line.find(key_words[4]) != -1:
                rods_data[2] = [float(A) for A in line[2:].replace('\n', ' ').replace('\t', ' ').split(' ') if A != '']

            elif line.find(key_words[5]) != -1:
                rods_data[3] = [float(E) for E in line[2:].replace('\n', ' ').replace('\t', ' ').split(' ') if E != '']

            elif line.find(key_words[6]) != -1:
                rods_data[4] = [float(Sigma) for Sigma in line[len(key_words[6]) + 1:].replace('\n', ' ').replace('\t', ' ').split(' ') if Sigma != '']

            elif line.find(key_words[7]) != -1:
                rods_forces = [float(rod_force) for rod_force in line[len(key_words[7]) + 1:].replace('\n', ' ').replace('\t', ' ').split(' ') if rod_force != '']

            elif line.find(key_words[8]) != -1:
                nodes_forces = [float(node_force) for node_force in line[len(key_words[8]) + 1:].replace('\n', ' ').replace('\t', ' ').split(' ') if node_force != '']

        file.close()

        if check_file() == True:
            break

        print()
        input(f"Исправьте данные в файле {file_name} и нажмите Enter для продолжения: ")

    print('Данные считаны!\n')
    return True


# Проверка значений из файла
def check_file():

    data_check = True

    key_words = ['L', 'A', 'E', 'Sigma']

    if n_rod < 1:

        data_check = False
        print(f"Ошибка! n_rod не может быть равно нулю или меньше нуля.")

    if hard_join_start not in [0, 1]:

        data_check = False
        print(f"Ошибка! Неверное значение в hard_join_start.")

    if hard_join_end not in [0, 1]:

        data_check = False
        print(f"Ошибка! Неверное значение в hard_join_end.")

    if (hard_join_start == 0) and (hard_join_end == 0):

        data_check = False
        print("Ошибка! Конструкция должна иметь хотя бы одну опору.")

    for i in range(1, len(key_words) + 1):

        if len(rods_data[i]) != n_rod:

            data_check = False
            print(f"Ошибка! Неверное количество элементов в {key_words[i - 1]}.")

    if len(rods_forces) != n_rod:

        data_check = False
        print(f"Ошибка! Неверное количество элементов в rods_forces.")

    if len(nodes_forces) != (n_rod + 1):

        data_check = False
        print(f"Ошибка! Неверное количество элементов в nodes_forces.")

    return data_check


# Запись в файл
def write_data():

    file_name = str(input('Введите название файла: '))

    # Проверка наличия расширения
    if file_name.find('.txt') == -1:
        file_name += '-project.txt'

    else:
        file_name = file_name.replace('.txt', '-project.txt')

    file = open(file_name, 'w')
    data_change = ['L\t', 'A\t', 'E\t', 'Sigma\t']

    file.write(f'N = {n_rod}\n')
    file.write(f'Hard_join_start = {hard_join_start}\n')
    file.write(f'Hard_join_end = {hard_join_end}\n')

    for i in range(len(data_change)):

        file.write(f'{data_change[i]}')

        for j in range(n_rod):

            file.write(f'{rods_data[i + 1][j]}\t')

        file.write('\n')

    file.write('rods_forces\t')
    for i in range(n_rod):

        file.write(f'{round(rods_forces[i], 3)}\t')

    file.write('\n')

    file.write('nodes_forces\t')
    for i in range(n_rod + 1):

        file.write(f'{round(nodes_forces[i], 3)}\t')

    file.write('\n')

    file.close()

    try:

        os.mkdir(f"{file_name[:-4]}-folder")

    except FileExistsError:

        print("Файл уже существует.")
        file_name = file_name.replace('-', '1-')
        os.rename(file_name.replace('1-', '-'), file_name)
        print(f"Название было изменено на '{file_name}'.")
        os.mkdir(f"{file_name[:-4]}-folder")

    shutil.move(f"{file_name}", f"{file_name[:-4]}-folder")

    print('Файл и папка проекта успешно созданы!')
    print(f'Название файла: {file_name}')
    print(f'Название папки: {file_name[:-4]}-folder\n')
    return file_name


# Создание примера файла
def create_example():

    choose = int(input('1 - Создать шаблон; 2 - Создать пример: '))

    if choose == 1:

        file = open('template.txt', 'w')

        file.write('N =\n')
        file.write('Hard_join_start = \n')
        file.write('Hard_join_end = \n')
        file.write('L\n')
        file.write('A\n')
        file.write('E\n')
        file.write('Sigma\n')
        file.write('rods_forces\n')
        file.write('nodes_forces\n')

        file.close()
        print('Шаблон создан!')
        print('Название template.txt\n')

    if choose == 2:

        file = open('example.txt', 'w')

        file.write('N = 2\n')
        file.write('Hard_join_start = 1\n')
        file.write('Hard_join_end = 1\n')
        file.write('L\t100\t120\n')
        file.write('A\t40\t45\n')
        file.write('E\t1\t2\n')
        file.write('Sigma\t3\t4\n')
        file.write('rods_forces\t100\t-20\n')
        file.write('nodes_forces\t100\t-30\t0\n')

        file.close()
        print('Пример создан!')
        print('Название example.txt\n')


# Изменение данных
def change_data(file_name):

    input(f'Внесите изменения в файл проекта ({file_name}), после чего нажмите Enter для продолжения: ')
    read_data(file_name)
    print('Данные успешно изменены!\n')


#------------------------------ Процессор ------------------------------

# Расчёт компонентов N_x, σ_x, U_x
def calculation_components():

    global matrix_A, vector_B, vector_Delta, U_x, N_x, Sigma_x
    global lengths

    n = n_rod + 1
    matrix_A = np.zeros((n, n))
    vector_B = np.zeros(n)
    lengths = rods_data[1]

    for i in range(n_rod):

        E = rods_data[3][i] 
        A = rods_data[2][i]  
        L = rods_data[1][i]  
        k = E * A / L
        matrix_A[i][i] += k
        matrix_A[i + 1][i + 1] += k
        matrix_A[i][i + 1] -= k
        matrix_A[i + 1][i] -= k

    if hard_join_start:

        matrix_A[0, :] = 0
        matrix_A[:, 0] = 0
        matrix_A[0, 0] = 1

    if hard_join_end:

        matrix_A[-1, :] = 0
        matrix_A[:, -1] = 0
        matrix_A[-1, -1] = 1

    for i in range(n_rod):

        F_longitudinal = rods_forces[i]  
        L = lengths[i]
        vector_B[i] += 0.5 * F_longitudinal * L
        vector_B[i + 1] += 0.5 * F_longitudinal * L

    for i, node_force in enumerate(nodes_forces):

        vector_B[i] += node_force

    if hard_join_start == True:

        vector_B[0] = 0

    if hard_join_end == True:

        vector_B[-1] = 0

    vector_Delta = np.linalg.solve(matrix_A, vector_B)

    N_x = []
    Sigma_x = []
    U_x = []

    for i in range(n_rod):

        E = rods_data[3][i]
        A = rods_data[2][i]
        L = lengths[i]
        U1, U2 = vector_Delta[i], vector_Delta[i + 1]
        N1 = ((E * A) / (L)) * (U2 - U1) + (rods_forces[i] * L / 2) 
        N2 = ((E * A) / (L)) * (U2 - U1) + ((rods_forces[i] * L / 2) * (1 - 2))  
        Sigma1 = N1 / A  
        Sigma2 = N2 / A  

        N_x.append([N1, N2])
        Sigma_x.append([Sigma1, Sigma2])
        U_x.append([U1, U2])

    write_res_to_file(matrix_A, vector_B, vector_Delta, N_x, Sigma_x, U_x)

    print("Процессор окончил работу.")
    input("Нажмите Enter, чтобы продолжить... ")
    print()


# Записываем результаты работы процессора в файл
def write_res_to_file(matrix_A, vector_B, vector_Delta, N_x, Sigma_x, vector_U):

    res_file = 'res-' + project_file_name

    with open(res_file, 'w') as file:

        file.write("Матрица жёсткости A:\n")
        for row in matrix_A:

            file.write("\t".join(map(lambda x: f"{round(x, 3)}", row)) + "\n")

        file.write("\n")

        file.write("Вектор нагрузок B:\n")
        file.write("\t".join(map(lambda x: f"{round(x, 3)}", vector_B)) + "\n\n")

        file.write("Вектор перемещений Delta:\n")
        file.write("\t".join(map(lambda x: f"{round(x, 3)}", vector_Delta)) + "\n\n")

        file.write("Усилия N_x:\n")
        for i, N in enumerate(N_x, start=1):

            file.write(f"Стержень {i}: {round(N[0], 3)}, {round(N[1], 3)}\n")

        file.write("\n")

        file.write("Напряжения Sigma_x:\n")
        for i, Sigma in enumerate(Sigma_x, start=1):

            file.write(f"Стержень {i}: {round(Sigma[0], 3)}, {round(Sigma[1], 3)}\n")

        file.write("\n")

        file.write("Перемещения U (по узлам):\n")
        for i, U in enumerate(vector_U, start=1):

            file.write(f"Стержень {i}: U1={round(U[0], 3)}, U2={round(U[1], 3)}\n")

    try:

        shutil.move(res_file, project_file_name.replace('.txt', '-folder'))

    except Exception:

        print(f"{res_file} уже существует!")
        print('Данные изменены!')
        os.remove(f"{project_file_name.replace('.txt', '-folder')}/{res_file}")
        shutil.move(res_file, project_file_name.replace('.txt', '-folder'))

    print(f"Результаты расчётов записаны в файл {res_file} в папке проекта.\n")


#---------------------------- Пост-процессор ----------------------------

# Отображение результатов в табличной форме в xlsx файле
def table_represent(): 

    wb = Workbook()
    ws = wb.active

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    while True:

        check = int(input(f"Выберите стержень для отображения значений в точке ({''.join(str(i) + ', ' for i in [j + 1 for j in range(n_rod - 1)])}{n_rod if n_rod >= 1 else ''})(0 - если стержень не выбран): "))

        if 0 <= check <= n_rod:

            break

        else:

            print("Введено неправильное значение! Попробуйте ещё раз.\n")

    points_components = []
    step = 0

    if check != 0:

        while True:

            point = float(input(f"Выберите точку на {check}-ом стержне (от 0 до {rods_data[1][check - 1]}): "))

            if 0 <= point <= rods_data[1][check - 1]:

                break

            else:

                print("Введено неправильное значение! Попробуйте ещё раз.\n")

    while True:

        step = float(input(f"Введите шаг для отображения компонентов стержней (0 - без шага)(от 0.1 до {min(rods_data[1]) / 2}): "))

        if 0 <= step <= (min(rods_data[1]) / 2):

            break

        else:

            print("Введено неправильное значение! Попробуйте ещё раз.\n")

    if step != 0:

        amount = len(list(j for j in np.arange(0, max(rods_data[1]), step)))

    temp = []

    if (step != 0):

        for i in range(n_rod):

            temp.clear()

            for j in np.arange(0, rods_data[1][i], step):

                x1 = 0
                x2 = rods_data[1][i]
                y1 = N_x[i][0]
                y2 = N_x[i][1]

                y_N = y1 + (j - x1) / (x2 - x1) * (y2 - y1)

                x1 = 0
                x2 = rods_data[1][i]
                y1 = Sigma_x[i][0]
                y2 = Sigma_x[i][1]

                y_Sigma = y1 + (j - x1) / (x2 - x1) * (y2 - y1)

                if rods_forces[i] != 0:

                    if i == 0:

                        C = 0

                    else:

                        C = U_x[i - 1][1]

                    x0 = (N_x[i][0] / rods_forces[i])

                    U_x[i].append(x0)
                    U_x[i].append(((N_x[i][0] * x0) + ((-rods_forces[i] * x0 ** 2) / 2)) / rods_data[2][i] + C)

                    coef1 = [0, U_x[i][0]]
                    coef2 = [rods_data[1][i], U_x[i][1]]
                    coef3 = [U_x[i][2], U_x[i][3]]

                    try:

                        a, b, c = find_parabola_coefficients([coef1, coef2, coef3])

                    except Exception:

                        print("Произошла непредвиденная ошибка. Попробуйте снова.\n")
                        return

                    y_U = a * j ** 2 + b * j + c

                else:

                    x1 = 0
                    x2 = rods_data[1][i]
                    y1 = U_x[i][0]
                    y2 = U_x[i][1]

                    y_U = y1 + (j - x1) / (x2 - x1) * (y2 - y1)

                temp.append([y_N, y_Sigma, y_U])

            points_components.append(temp.copy())

    center_alignment = Alignment(horizontal="center", vertical="center")

    if step == 0:

        ws.append(["Номер стержня", "Продольная сила в начале", "Продольная сила в конце", "Нормальное напряжение в начале", "Нормальное напряжение в конце", "Перемещение в начале", "Перемещения в конце"])

    else:

        headers = ["Номер стержня"]
        headers.extend([f"Продольная сила в точке {round(k, 3)}" for k in np.arange(0, max(rods_data[1]), step)])
        headers.extend([f"Нормальное напряжение в точке {round(k, 3)}" for k in np.arange(0, max(rods_data[1]), step)])
        headers.extend([f"Перемещения в точке {round(k, 3)}" for k in np.arange(0, max(rods_data[1]), step)])

        ws.append(headers)

    for i in range(n_rod):

            if step == 0:

                ws.append([

                    f"{i + 1}",
                    f"{round(N_x[i][0], 3)}",
                    f"{round(N_x[i][1], 3)}",
                    f"{round(Sigma_x[i][0], 3)}",
                    f"{round(Sigma_x[i][1], 3)}",
                    f"{round(U_x[i][0], 3)}",
                    f"{round(U_x[i][1], 3)}"

                ])

            else:

                values = [f"{i + 1}"]
                values.extend([round(points_components[i][k][0], 3) if k < (rods_data[1][i] / step) else '-' for k in range(amount)])
                values.extend([round(points_components[i][k][1], 3) if k < (rods_data[1][i] / step) else '-' for k in range(amount)])
                values.extend([round(points_components[i][k][2], 3) if k < (rods_data[1][i] / step) else '-' for k in range(amount)])

                ws.append(values)

    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):

        for cell in row:

            cell.border = thin_border
            cell.alignment = center_alignment

    if check != 0:

        temp = []

        x1 = 0
        x2 = rods_data[1][(check - 1)]
        y1 = N_x[(check - 1)][0]
        y2 = N_x[(check - 1)][1]

        y = y1 + (point - x1) / (x2 - x1) * (y2 - y1)

        temp.append(y)

        x1 = 0
        x2 = rods_data[1][(check - 1)]
        y1 = Sigma_x[(check - 1)][0]
        y2 = Sigma_x[(check - 1)][1]

        y = y1 + (point - x1) / (x2 - x1) * (y2 - y1)

        temp.append(y)

        if rods_forces[(check - 1)] != 0:

            if (check - 1) == 0:

                C = 0

            else:

                C = U_x[check - 2][1]

            x0 = (N_x[(check - 1)][0] / rods_forces[(check - 1)])

            U_x[(check - 1)].append(x0)
            U_x[(check - 1)].append(((N_x[(check - 1)][0] * x0) + ((-rods_forces[(check - 1)] * x0 ** 2) / 2)) / rods_data[2][check - 1] + C)

            coef1 = [0, U_x[(check - 1)][0]]
            coef2 = [rods_data[1][(check - 1)], U_x[(check - 1)][1]]
            coef3 = [U_x[(check - 1)][2], U_x[(check - 1)][3]]

            a, b, c = find_parabola_coefficients([coef1, coef2, coef3])

            y = a * point ** 2 + b * point + c

        else:

            x1 = 0
            x2 = rods_data[1][(check - 1)]
            y1 = U_x[(check - 1)][0]
            y2 = U_x[(check - 1)][1]

            y = y1 + (point - x1) / (x2 - x1) * (y2 - y1)

        temp.append(y)

        head = ["Номер стержня", "Точка на стержне", "Продольная сила в точке", "Нормальное напряжение в точке", "Перемещение в точке"]
        point_value = [str(check), str(point), round(temp[0], 3), round(temp[1], 3), round(temp[2], 3)]

        m_row_1 = ws.max_row + 3
        m_row_2 = ws.max_row + 4

        for i in range(len(head)):

            current_cell_1 = ws.cell(row=m_row_1, column=1 + i)
            current_cell_2 = ws.cell(row=m_row_2, column=1 + i)

            current_cell_1.value = head[i]
            current_cell_2.value = point_value[i]

        for row in ws.iter_rows(min_row=m_row_1, max_row=m_row_2, min_col=1, max_col=len(head)):

            for cell in row:

                cell.border = thin_border
                cell.alignment = center_alignment

    xlsx_file = project_file_name.replace(".txt", "-components.xlsx")
    wb.save(xlsx_file)

    try:

        shutil.move(xlsx_file, project_file_name.replace('.txt', '-folder'))

    except Exception:

        print(f"{xlsx_file} уже существует!")

        try:

            os.remove(f"{project_file_name.replace('.txt', '-folder')}/{xlsx_file}")

        except Exception:

            print('Файл используется другим процессом, закройте файл и попробуйте ещё раз.')
            print('Данные не заменены!\n')
            return

        shutil.move(xlsx_file, project_file_name.replace('.txt', '-folder'))
        print('Данные заменены!')

    print(f"Результаты записаны в файл {xlsx_file} в папке проекта.\n")


# Отображение нормальных напряжений в xlsx файле
def normal_tensions():

    wb = Workbook()
    ws = wb.active

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    center_alignment = Alignment(horizontal="center", vertical="center")

    ws.append(["Номер стержня", "Нормальное напряжение в начале", "Нормальное напряжение в конце", "Допускаемое напряжение"])

    for i in range(n_rod):

        ws.append([f"{i + 1}", f"{round(Sigma_x[i][0], 3)}", f"{round(Sigma_x[i][1], 3)}", rods_data[4][i]])

        start_cell = ws.cell(row=ws.max_row, column=2)
        end_cell = ws.cell(row=ws.max_row, column=3)

        if abs(float(start_cell.value)) > rods_data[4][i]:

            start_cell.font = Font(color="FF0000")

        else:

            start_cell.font = Font(color="00FF00")

        if abs(float(end_cell.value)) > rods_data[4][i]:

            end_cell.font = Font(color="FF0000")

        else:

            end_cell.font = Font(color="00FF00")

        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=1, max_col=ws.max_column):

            for cell in row:

                cell.border = thin_border
                cell.alignment = center_alignment

    xlsx_file = project_file_name.replace(".txt", "-tensions.xlsx")
    wb.save(xlsx_file)

    try:

        shutil.move(xlsx_file, project_file_name.replace('.txt', '-folder'))

    except Exception:

        print(f"{xlsx_file} уже существует!")

        try:

            os.remove(f"{project_file_name.replace('.txt', '-folder')}/{xlsx_file}")

        except Exception:

            print('Файл используется другим процессом, закройте файл и попробуйте ещё раз.')
            print('Данные не заменены!\n')
            return

        shutil.move(xlsx_file, project_file_name.replace('.txt', '-folder'))
        print('Данные заменены!')

    print(f"Результаты записаны в файл {xlsx_file} в папке проекта.\n")


# Подсчёт коэффициентов параболы
def find_parabola_coefficients(points):

    x1, y1 = points[0]
    x2, y2 = points[1]
    x3, y3 = points[2]

    A = np.array([
        [x1 ** 2, x1, 1],
        [x2 ** 2, x2, 1],
        [x3 ** 2, x3, 1]
    ])

    B = np.array([y1, y2, y3])
    a, b, c = np.linalg.solve(A, B)

    return a, b, c


# Отрисовка эпюр
def draw_epure(canvas, type):

    start = mid_height
    end_x = start
    max_height = max(rods_data[2])
    epure_limit = 130

    epure_line = middle + (max_height * 2.3) + 160

    for i in range(n_rod):

        canvas.create_line(end_x, middle, end_x, epure_line + 150, width=2)
        end_x += rods_data[1][i]
    
    canvas.create_line(end_x, middle, end_x, epure_line + 150, width=2)
    canvas.create_line(start, epure_line, end_x, epure_line, width=3)
    end_x = start

    if type == 4:  

        canvas.create_text(mid_height - 30, epure_line - 15, text=f"N_x", font=("Times New Roman", 18, "bold"))
        canvas.create_text(mid_height - 30, epure_line + 15, text=f"qL", font=("Times New Roman", 10, "bold"))

        res = []

        for i in range(n_rod):

            res.append(N_x[i][0] * 10)
            res.append(N_x[i][1] * 10)

        conv = epure_limit / max(res)
        res.clear()

        for i in range(n_rod):

            canvas.create_text(end_x + 20, (epure_line - (N_x[i][0] * 10) if conv > 1 else epure_line - (N_x[i][0] * 10 * conv / 1.5)) - 20, text=f"{round(abs(N_x[i][0]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')
            canvas.create_line(end_x, epure_line - (N_x[i][0] * 10) if conv > 1 else epure_line - (N_x[i][0] * 10 * conv / 1.5), end_x + (rods_data[1][i]), epure_line - (N_x[i][1] * 10) if conv > 1 else epure_line - (N_x[i][1] * 10 * conv / 1.5), width=4, fill='#172ba7')

            for j in np.arange(end_x, end_x + (rods_data[1][i]), 10):

                x1 = end_x
                x2 = end_x + (rods_data[1][i])
                y1 = epure_line - (N_x[i][0] * 10) if conv > 1 else epure_line - (N_x[i][0] * 10 * conv / 1.5)
                y2 = epure_line - (N_x[i][1] * 10) if conv > 1 else epure_line - (N_x[i][1] * 10 * conv / 1.5)

                y = y1 + (j - x1) / (x2 - x1) * (y2 - y1)

                canvas.create_line(j, epure_line, j, y, width=2, fill='#172ba7')

            end_x += rods_data[1][i]

            canvas.create_text(end_x - 20, (epure_line - (N_x[i][1] * 10) if conv > 1 else epure_line - (N_x[i][1] * 10 * conv / 1.5)) + 20, text=f"{round(abs(N_x[i][1]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')

    elif type == 5:

        canvas.create_text(mid_height - 30, epure_line - 15, text=f"σ_x", font=("Times New Roman", 18, "bold"))
        canvas.create_text(mid_height - 30, epure_line + 15, text=f"qL / A", font=("Times New Roman", 10, "bold"))

        res = []

        for i in range(n_rod):

            res.append(Sigma_x[i][0] * 10)
            res.append(Sigma_x[i][1] * 10)

        conv = epure_limit / max(res)
        res.clear()

        for i in range(n_rod):

            for j in np.arange(end_x, end_x + (rods_data[1][i]), 10):

                x1 = end_x
                x2 = end_x + (rods_data[1][i])
                y1 = epure_line - (Sigma_x[i][0] * 10) if conv > 1 else epure_line - (Sigma_x[i][0] * 10 * conv / 1.5)
                y2 = epure_line - (Sigma_x[i][1] * 10) if conv > 1 else epure_line - (Sigma_x[i][1] * 10 * conv / 1.5)

                y = y1 + (j - x1) / (x2 - x1) * (y2 - y1)

                canvas.create_line(j, epure_line, j, y, width=2, fill='#172ba7')

            canvas.create_text(end_x + 20, (epure_line - (Sigma_x[i][0] * 10) if conv > 1 else epure_line - (Sigma_x[i][0] * 10 * conv / 1.5)) - 20, text=f"{round(abs(Sigma_x[i][0]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')
            canvas.create_line(end_x, epure_line - (Sigma_x[i][0] * 10) if conv > 1 else epure_line - (Sigma_x[i][0] * 10 * conv / 1.5), end_x + (rods_data[1][i]), epure_line - (Sigma_x[i][1] * 10) if conv > 1 else epure_line - (Sigma_x[i][1] * 10 * conv / 1.5), width=4, fill='#172ba7')

            end_x += rods_data[1][i]

            canvas.create_text(end_x - 20, (epure_line - (Sigma_x[i][1] * 10) if conv > 1 else epure_line - (Sigma_x[i][1] * 10 * conv / 1.5)) + 20, text=f"{round(abs(Sigma_x[i][1]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')

    elif type == 6:

        canvas.create_text(mid_height - 30, epure_line - 15, text=f"U_x", font=("Times New Roman", 18, "bold"))
        canvas.create_text(mid_height - 30, epure_line + 15, text=f"qL^2 / EA", font=("Times New Roman", 10, "bold"))

        res = []

        for i in range(n_rod):

            if rods_forces[i] != 0:

                if i == 0:

                    C = 0

                else:

                    C = U_x[i - 1][1]

                x0 = (N_x[i][0] / (rods_forces[i] / 1000))

                U_x[i].append(x0)
                U_x[i].append(((N_x[i][0] * x0) + ((-(rods_forces[i] / 1000) * x0 ** 2) / 2)) / (rods_data[2][i] / 3) + C)

                coef1 = [0, U_x[i][0]]
                coef2 = [rods_start[i], U_x[i][1]]
                coef3 = [U_x[i][2], U_x[i][3]]

                a, b, c = find_parabola_coefficients([coef1, coef2, coef3])

                temp = []

                step = 0.1 
                for x in np.arange(0, rods_start[i], step):

                    y = a * x ** 2 + b * x + c
                    y_next = a * (x + step) ** 2 + b * (x + step) + c

                    res.append(y * 10)

            else: 

                res.append(U_x[i][0] * 10)
                res.append(U_x[i][1] * 10)

        conv = epure_limit / max(res)
        res.clear()

        for i in range(n_rod):

            canvas.create_text(end_x + 20, (epure_line - (U_x[i][0] * 10) if conv > 1 else epure_line - (U_x[i][0] * 10 * conv / 1.5)) - 20, text=f"{round(abs(U_x[i][0]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')

            if rods_forces[i] != 0:

                if i == 0:

                    C = 0

                else:

                    C = U_x[i - 1][1]

                x0 = (N_x[i][0] / (rods_forces[i] / 1000))

                U_x[i].append(x0)
                U_x[i].append(((N_x[i][0] * x0) + ((-(rods_forces[i] / 1000) * x0 ** 2) / 2)) / (rods_data[2][i] / 3) + C)

                coef1 = [0, U_x[i][0]]
                coef2 = [rods_start[i], U_x[i][1]]
                coef3 = [U_x[i][2], U_x[i][3]]

                a, b, c = find_parabola_coefficients([coef1, coef2, coef3])

                temp = []

                step = 0.1 
                for x in np.arange(0, rods_start[i], step):

                    y = a * x ** 2 + b * x + c
                    y_next = a * (x + step) ** 2 + b * (x + step) + c

                    temp.append([x, y])

                    canvas.create_line(end_x + (x * 100 / length_conv), epure_line, end_x + (x * 100 / length_conv), epure_line - (y * 10) if conv > 1 else epure_line - (y * 10 * conv / 1.5), width=2, fill='#172ba7')
                    canvas.create_line(end_x + (x * 100 / length_conv), epure_line - (y * 10) if conv > 1 else epure_line - (y * 10 * conv / 1.5), end_x + (x + step) * 100 / length_conv, epure_line - (y_next * 10) if conv > 1 else epure_line - (y_next * 10 * conv / 1.5), fill="#172ba7", width=4)

                if 0 < U_x[i][2] < rods_start[i]:

                    canvas.create_text(end_x + U_x[i][2] * L_si / length_conv, (epure_line - (U_x[i][3] * 10) if conv > 1 else epure_line - (U_x[i][3] * 10 * conv / 1.5)) + 10, text=f"{round(abs(U_x[i][3]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')

            else:

                canvas.create_line(end_x, epure_line - (U_x[i][0] * 10) if conv > 1 else epure_line - (U_x[i][0] * 10 * conv / 1.5), end_x + (rods_data[1][i]), epure_line - (U_x[i][1] * 10) if conv > 1 else epure_line - (U_x[i][1] * 10 * conv / 1.5), width=4, fill='#172ba7')

                for j in np.arange(end_x, end_x + (rods_data[1][i]), 10):

                    x1 = end_x
                    x2 = end_x + (rods_data[1][i])
                    y1 = epure_line - (U_x[i][0] * 10) if conv > 1 else epure_line - (U_x[i][0] * 10 * conv / 1.5)
                    y2 = epure_line - (U_x[i][1] * 10) if conv > 1 else epure_line - (U_x[i][1] * 10 * conv / 1.5)
                    
                    y = y1 + (j - x1) / (x2 - x1) * (y2 - y1)
                
                    canvas.create_line(j, epure_line, j, y, width=2, fill='#172ba7')
                
            end_x += rods_data[1][i]

            canvas.create_text(end_x - 20, (epure_line - (U_x[i][1] * 10) if conv > 1 else epure_line - (U_x[i][0] * 10 * conv / 1.5)) + 20, text=f"{round(abs(U_x[i][1]), 3)}", font=("Times New Roman", 14, "bold"), fill='#107c3e')


#-------------------------------- Отчёт --------------------------------

# Функция для формирования отчёта
def report_generation():

    author = input("Введите имя автора отчёта: ")
    role = input("Введите группу: ")

    while True:

        while True:

            if (choose_rep := int(input("Выберите формат отчёта:\n1) Word\n2) Html\n3) Выход из генератора отчётов\n-> "))) in [1, 2, 3]:

                break

            else:

                print("Введено неверное значение! Попробуйте ещё раз.\n")

        if choose_rep == 1:

            doc_name = "report-" + project_file_name[:project_file_name.find('.txt')] + ".docx"
            doc = docx.Document()

            stankin_image = "stankin.png"
            stankin_p = doc.add_paragraph()
            stankin_run = stankin_p.add_run()
            stankin_run.add_picture(stankin_image)
            stankin_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("\n\n\n")

            title = doc.add_heading("Отчёт", level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
            for run in title.runs:

                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(24)

            doc.add_paragraph("\n\n\n\n\n")

            role_p = doc.add_paragraph(role)
            role_p.bold = True
            role_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in role_p.runs:

                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(16)

            author_p = doc.add_paragraph(author + "\t\t\t\t     ___________________(Подпись)")
            author_p.bold = True
            author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in author_p.runs:

                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(16)

            date_p = doc.add_paragraph(str(datetime.now().date()))
            date_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in date_p.runs:

                run.font.size = Pt(16)

            doc.add_paragraph("\n\n\n\n\n\n\n\n\n\n")

            moscow = doc.add_paragraph("Москва " + str(datetime.now().year))
            moscow.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in moscow.runs:

                run.font.size = Pt(14)

            doc.add_page_break()

            parametrs = doc.add_paragraph("Параметры стержневой системы")
            parametrs.bold = True
            parametrs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in parametrs.runs:

                run.font.size = Pt(18)

            doc.add_paragraph("\n")

            with open(project_file_name.replace('.txt', '-folder') + "/" + project_file_name, 'r') as file:

                for line in file:

                    doc.add_paragraph(str(line))

            doc.add_page_break()

            results = doc.add_paragraph("Результаты работы процессора")
            results.bold = True
            results.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in results.runs:

                run.font.size = Pt(18)

            doc.add_paragraph("\n")

            with open(project_file_name.replace('.txt', '-folder') + '/res-' + project_file_name, 'r') as file:

                for line in file:

                    doc.add_paragraph(str(line))

            doc.add_page_break()

            tables = doc.add_paragraph("Таблицы")
            tables.bold = True
            tables.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in tables.runs:

                run.font.size = Pt(18)

            doc.add_paragraph("\n")

            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'  

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Номер стержня"
            hdr_cells[1].text = "Продольная сила в начале"
            hdr_cells[2].text = "Продольная сила в конце"
            hdr_cells[3].text = "Нормальное напряжение в начале"
            hdr_cells[4].text = "Нормальное напряжение в конце"
            hdr_cells[5].text = "Перемещение в начале"
            hdr_cells[6].text = "Перемещение в конце"

            for i in range(n_rod):

                row_cells = table.add_row().cells
                row_cells[0].text = f"{i + 1}"
                row_cells[1].text = str(round(N_x[i][0], 3))
                row_cells[2].text = str(round(N_x[i][1], 3))
                row_cells[3].text = str(round(Sigma_x[i][0], 3))
                row_cells[4].text = str(round(Sigma_x[i][1], 3))
                row_cells[5].text = str(round(U_x[i][0], 3))
                row_cells[6].text = str(round(U_x[i][1], 3))

            for row in table.rows:

                for cell in row.cells:

                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph("\n\n")

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Номер стержня"
            hdr_cells[1].text = "Напряжение в начале"
            hdr_cells[2].text = "Напряжение в конце"
            hdr_cells[3].text = "Допускаемое напряжение"

            for i in range(n_rod):

                row_cells = table.add_row().cells

                row_cells[0].text = f"{i + 1}"
                row_cells[1].text = str(round(Sigma_x[i][0], 3))
                row_cells[2].text = str(round(Sigma_x[i][1], 3))
                row_cells[3].text = str(rods_data[4][i])

                if abs(Sigma_x[i][0]) > rods_data[4][i]:

                    row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                else:

                    row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 255, 0) 

                if abs(Sigma_x[i][1]) > rods_data[4][i]:

                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                else:

                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 255, 0) 

                for cell in row_cells:

                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.save(doc_name)

            try:

                shutil.move(doc_name, project_file_name.replace('.txt', '-folder'))

            except Exception:

                print(f"{doc_name} уже существует!")

                try:

                    os.remove(f"{project_file_name.replace('.txt', '-folder')}/{doc_name}")

                except Exception:

                    print('Файл используется другим процессом, закройте файл и попробуйте ещё раз.')
                    print('Данные не заменены!\n')
                    return

                shutil.move(doc_name, project_file_name.replace('.txt', '-folder'))
                print('Данные заменены!')

            print(f"Отчёт записан в файл {doc_name} в папке проекта.\n")

            global word_created
            word_created = True


        elif choose_rep == 2:

            html_file_name = "report-" + project_file_name[:project_file_name.find('.txt')] + ".html"

            html = f"""
            <!DOCTYPE html>
            <html lang='ru'>
            <head>
                <meta charset='UTF-8'>
                <meta name='viewport' content='width=device-width, initial-scale=1.0'>
                <title>Отчёт</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 0 20px; }}
                    h1, h2, h3, h4 {{ text-align: center; color: black; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    table, th, td {{ border: 1px solid black; }}
                    th, td {{ padding: 8px; text-align: center; }}
                    .green {{ color: green; }}
                    .red {{ color: red; }}
                    .signature {{ text-align: center; margin-top: 50px; }}
                    .date, .location {{ text-align: right; }}
                </style>
            </head>
            <body>
                <img src='stankin.png' alt='Logo' style='display: block; margin: 0 auto;'>
                <h1>Отчёт</h1>

                <h3>{role}</h3>
                <div class='signature'>
                    <p>{author}<span style='float: right;'>___________________(Подпись)</span></p>
                </div>
                <p class='date'>{datetime.now().date()}</p>
                <p class='location'>Москва {datetime.now().year}</p>

                <hr>

                <h2>Параметры стержневой системы</h2>
                <pre>
            """

            with open(project_file_name.replace('.txt', '-folder') + "/" + project_file_name, 'r') as file:

                for line in file:

                    html += line

            html += """</pre>

                <hr>

                <h2>Результаты работы процессора</h2>
                <pre>
            """

            with open(project_file_name.replace('.txt', '-folder') + '/res-' + project_file_name, 'r') as file:

                for line in file:

                    html += line

            html += """</pre>

                <hr>

                <h2>Таблицы</h2>

                <h3>Основные параметры</h3>
                <table>
                    <tr>
                        <th>Номер стержня</th>
                        <th>Продольная сила в начале</th>
                        <th>Продольная сила в конце</th>
                        <th>Нормальное напряжение в начале</th>
                        <th>Нормальное напряжение в конце</th>
                        <th>Перемещение в начале</th>
                        <th>Перемещение в конце</th>
                    </tr>
            """

            for i in range(n_rod):

                html += f"""
                    <tr>
                        <td>{i + 1}</td>
                        <td>{round(N_x[i][0], 3)}</td>
                        <td>{round(N_x[i][1], 3)}</td>
                        <td>{round(Sigma_x[i][0], 3)}</td>
                        <td>{round(Sigma_x[i][1], 3)}</td>
                        <td>{round(U_x[i][0], 3)}</td>
                        <td>{round(U_x[i][1], 3)}</td>
                    </tr>
                """

            html += """</table>

                <h3>Напряжения и допускаемые значения</h3>
                <table>
                    <tr>
                        <th>Номер стержня</th>
                        <th>Напряжение в начале</th>
                        <th>Напряжение в конце</th>
                        <th>Допускаемое напряжение</th>
                    </tr>
            """

            for i in range(n_rod):

                sigma_start_class = "red" if abs(Sigma_x[i][0]) > rods_data[4][i] else "green"
                sigma_end_class = "red" if abs(Sigma_x[i][1]) > rods_data[4][i] else "green"

                html += f"""
                    <tr>
                        <td>{i + 1}</td>
                        <td class='{sigma_start_class}'>{round(Sigma_x[i][0], 3)}</td>
                        <td class='{sigma_end_class}'>{round(Sigma_x[i][1], 3)}</td>
                        <td>{rods_data[4][i]}</td>
                    </tr>
                """

            html += """</table>

            </body>
            </html>
            """

            with open(html_file_name, 'w', encoding='utf-8') as file:

                file.write(html)

            try:

                shutil.move(html_file_name, project_file_name.replace('.txt', '-folder'))

            except Exception:

                print(f"{html_file_name} уже существует!")

                try:

                    os.remove(f"{project_file_name.replace('.txt', '-folder')}/{html_file_name}")

                except Exception:

                    print('Файл используется другим процессом, закройте файл и попробуйте ещё раз.')
                    print('Данные не заменены!\n')
                    return

                shutil.move(html_file_name, project_file_name.replace('.txt', '-folder'))
                print('Данные заменены!')

            print(f"Отчёт записан в файл {html_file_name} в папке проекта.\n")


        elif choose_rep == 3:

            return

#------------------------------------------------------------------------


if __name__ == "__main__":

    print('Добро пожаловать в САПР')
    print('© Кужба Полина, ИДБ-22-12\n\n')

    global project_file_name

    data_exist = False
    project_file_exist = False
    project_file_name = ''
    result_file_exist = False
    report_file_exist = False

    while True:

        print("Выберите действие:")
        print("1) Препроцессор")
        print("2) Процессор")
        print("3) Постпроцессор")
        print("4) Сформировать отчёт")
        print("5) Выход из программы")
        print("-> ", end='')

        while not validate_data(choose_proc := int(input())):
                    pass

        print()

        if choose_proc == 1:
            
            while True:

                print("Выберите действие:")
                print("1) Внесение данных")
                print("2) Изменение данных")
                print("3) Отрисовка конструкции")
                print("4) Отрисовка нагрузок")
                print("5) Сформировать файл проекта")
                print("6) Справка")
                print("7) Выход из препроцессора")
                print("-> ", end='')

                while not validate_data(choose := int(input())):
                    pass

                print()

                if choose == 1:

                    while True:

                        if data_exist == True:
                            
                            print('В программу уже внесены данные. Хотите их изменить? (1 - Да; 0 - Нет)-> ', end=' ')
                            if (int(input()) == 0):
                                break
                            
                        data_exist = False
                            
                        print('Выберите способ ввода данных:')
                        print('1) Из файла')
                        print('2) Из консоли')
                        print('3) Открыть проект')
                        print('4) Создать шаблон или пример файла')
                        print('5) Вернуться назад')
                        print("-> ", end='')

                        while not validate_data(choose_input := int(input())):
                            pass

                        print()

                        if choose_input == 1:

                            if read_data() == True:
                            
                                data_exist = True
                                
                            break

                        elif choose_input == 2:

                            input_data()
                            data_exist = True
                            break

                        elif choose_input == 3:

                            project_file_name = input("Введите название папки проекта или самого проекта: ")

                            if (project_file_name.find("-project") == -1):

                                project_file_name = project_file_name + "-project.txt"

                            elif (project_file_name.find("-folder") != -1):

                                project_file_name = project_file_name[:project_file_name.find("-folder")] + ".txt"

                            else:

                                project_file_name += ".txt"

                            project_file_exist = True

                            if read_data(project_file_name) == False:

                                print('Проект не был открыт! Попробуйте заново.\n')

                            else:

                                data_exist = True
                                print("Папка проекта открыта.\n")

                            break

                        elif choose_input == 4:
                            create_example()

                        elif choose_input == 5:
                            break

                        else:
                            print('Введено неправильное значение! Попробуйте снова.\n')

                elif (choose == 2) and (data_exist == True):

                    if not project_file_exist:
                        print('Не был создан файл проекта!')
                        print('Идёт создание...\n')
                        project_file_name = write_data()
                        project_file_exist = True

                    change_data(project_file_name)

                elif (choose == 3) and (data_exist == True):
                    draw_construction(1)

                elif (choose == 4) and (data_exist == True):
                    
                    while True:

                        print('Какие нагрузки отобразить:')
                        print('1) Нагрузки на стержни')
                        print('2) Нагрузки на узлы')
                        print('3) Вернуться назад')
                        print('-> ', end='')

                        while not validate_data(choose_forces := int(input())):
                            pass

                        print()

                        if choose_forces == 1:
                            draw_construction(2)

                        elif choose_forces == 2:
                            draw_construction(3)

                        elif choose_forces == 3:
                            break

                        else:
                            print('Введено неправильное значение! Попробуйте заново.\n')

                elif (choose == 5) and (data_exist == True):
                    project_file_name = write_data()
                    project_file_exist = True

                elif choose == 6:
                    print('Единицы измерения параметров:')
                    print("L[м] - метры")
                    print("A[М^2] - квадратные метры")
                    print("E[Н/м^2] - Ньютон на метр квадратный")
                    print("σ[МПа] - Мегапаскаль")
                    print("q[Н] - Ньютон на метр")
                    print("F[Н] - Ньютон\n")

                elif choose == 7:
                    break

                else:
                    print('Введено неправильное значение или данных нет! Попробуйте снова.\n')

        elif (choose_proc == 2) and (project_file_exist == True):

            calculation_components()
            result_file_exist = True

        elif (choose_proc == 3) and (result_file_exist == True):
            
            while True:

                print('Выберите действие:')
                print('1) Отображение результатов в табличной форме')
                print('2) Отображение нормальных напряжений в табличной форме')
                print('3) Отображение результатов в виде эпюры')
                print('4) Выход из постпроцессора')
                print('-> ', end='')

                choose_post = int(input())

                print()

                if choose_post == 1:

                    table_represent()

                elif choose_post == 2:

                    normal_tensions()

                elif choose_post == 3:

                    while True:

                        print('Какой компонент необходимо отобразить?')
                        print('1) N_x')
                        print('2) σ_x')
                        print('3) U_x')
                        print('4) Вернуться назад')
                        print('-> ', end='')

                        choose_epure = int(input())

                        print()

                        if choose_epure in [1, 2, 3]:

                            draw_construction(choose_epure + 3)

                        elif choose_epure == 4:
                            break

                        else:
                            print('Введено неправильное значение! Попробуйте снова.\n')

                elif choose_post == 4:
                    break

                else:
                    print('Введено неправильное значение! Попробуйте снова.\n')

        elif choose_proc == 4:

            if (result_file_exist == False):

                print("Необходимо сформировать файл проекта и файл результатов для того, чтобы сформировать отчёт.\n")

            else:

                report_generation()
                report_file_exist = True

        elif choose_proc == 5:

            print("Проверьте сохранили ли вы все необходимые данные:")
            print("1) Файл проекта - " + ("СОЗДАН" if project_file_exist else "НЕ СОЗДАН"))
            print("2) Файл результатов - " + ("СОЗДАН" if result_file_exist else "НЕ СОЗДАН"))
            print("3) Отчёт - " + ("СОЗДАН" if report_file_exist else "НЕ СОЗДАН"))
            print("")

            if (exit_app := int(input("Выйти из программы? (1 - Да; 0 - Нет): "))) == 1:

                print('Программа завершила работу.')
                sys.exit(1)

            print("")

        else:
            print('Введено неправильное значение! Попробуйте снова.\n')
